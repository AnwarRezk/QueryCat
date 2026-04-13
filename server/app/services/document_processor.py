"""
Unified document processing service — parses PDF, DOCX, and Markdown files
into LangChain Document chunks with session-scoped metadata.

Supported formats:
  .pdf  — PyPDFLoader (page-by-page extraction)
  .docx — python-docx (paragraphs + tables as structured text)
  .md   — Plain UTF-8 text with heading-aware splitting

Every chunk is stamped with:
  - session_id  : ties the chunk to a specific conversation
  - source      : original filename
  - upload_id   : unique ID for this upload batch
  - chunk_index : position in the chunk sequence
  - file_type   : 'pdf' | 'docx' | 'md'
"""
import logging
import os
from pathlib import Path
from typing import Literal

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings

logger = logging.getLogger(__name__)

FileType = Literal["pdf", "docx", "md"]

ALLOWED_EXTENSIONS: set[str] = {".pdf", ".docx", ".md"}


def _get_file_type(filename: str) -> FileType:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    elif ext == ".docx":
        return "docx"
    elif ext == ".md":
        return "md"
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )


def _load_pdf(file_path: str) -> list[Document]:
    """Load a PDF file page by page using PyPDFLoader."""
    loader = PyPDFLoader(file_path)
    return loader.load()


def _load_docx(file_path: str, filename: str) -> list[Document]:
    """
    Extract text from a .docx file using python-docx.

    Strategy:
    - Extract paragraphs as body text
    - Extract tables as pipe-delimited text rows
    - Return as a single Document (splitter handles chunking)
    """
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError:
        raise RuntimeError(
            "python-docx is required for Word document support. "
            "Run: uv pip install python-docx"
        )

    doc = DocxDocument(file_path)
    sections: list[str] = []

    # Extract paragraphs, preserving heading hierarchy
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style:
            sections.append(f"\n## {text}\n")
        elif "Title" in style:
            sections.append(f"\n# {text}\n")
        else:
            sections.append(text)

    # Extract tables as readable text
    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append(" | ".join(cells))
        if table_lines:
            sections.append("\n" + "\n".join(table_lines) + "\n")

    full_text = "\n".join(sections)
    logger.info("Extracted %d characters from DOCX: %s", len(full_text), filename)

    return [Document(page_content=full_text, metadata={"source": filename, "page": 0})]


def _load_markdown(file_path: str, filename: str) -> list[Document]:
    """
    Load a Markdown file as a LangChain Document.
    The text splitter will split on heading boundaries.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    logger.info("Loaded %d characters from Markdown: %s", len(content), filename)
    return [Document(page_content=content, metadata={"source": filename, "page": 0})]


async def process_document(
    file_path: str,
    filename: str,
    upload_id: str,
    session_id: str,
) -> list[Document]:
    """
    Load, split, and enrich a document for vector storage.

    Args:
        file_path:  Absolute path to the saved file on disk.
        filename:   Original filename (used in metadata + citations).
        upload_id:  Unique ID for this upload batch.
        session_id: Chat session this document belongs to.

    Returns:
        List of enriched Document chunks ready for embedding.
    """
    settings = get_settings()
    file_type = _get_file_type(filename)

    logger.info("Processing %s file: %s (session=%s)", file_type, filename, session_id)

    # ── Load raw pages/content based on file type ─────────────────────────────
    if file_type == "pdf":
        pages = _load_pdf(file_path)
    elif file_type == "docx":
        pages = _load_docx(file_path, filename)
    else:  # md
        pages = _load_markdown(file_path, filename)

    # ── Split into chunks ─────────────────────────────────────────────────────
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["## ", "\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(pages)

    # ── Enrich metadata on every chunk ────────────────────────────────────────
    for idx, chunk in enumerate(chunks):
        chunk.metadata.update(
            {
                "source": filename,
                "upload_id": upload_id,
                "session_id": session_id,   # ← KEY: scopes retrieval to this chat
                "file_type": file_type,
                "chunk_index": idx,
                "total_chunks": len(chunks),
            }
        )

    logger.info(
        "Processed '%s' → %d chunks (type=%s, session=%s)",
        filename,
        len(chunks),
        file_type,
        session_id,
    )
    return chunks
